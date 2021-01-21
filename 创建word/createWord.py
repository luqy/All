# -*- coding: utf-8 -*-
# @Time    : 2020/8/14 13:25
# @Author  : luqy
# @Email   :
# @File    : createWord.py
# @Software: PyCharm
from docx import Document
from docx.shared import Pt,RGBColor,Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches
document=Document()
# #document.add_heading('全网云功能测试报告',0) #表题，有下划线
# document.styles['Normal'].font.name = u'Helvetica'      #可换成word里面任意字体
# document.styles['Normal']._element.rPr.rFonts.set(qn('w:Helvetica'), 'Helvetica')
# p = document.add_paragraph()
# p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER    #段落文字居中设置
# run = p.add_run(u'全网云功能测试报告')
# run.bold = True
# run.font.color.rgb = RGBColor(0,0,0)             #颜色设置，这里是用RGB颜色
# run.font.size = Pt(24)
# document.add_picture(r'F:\untitled\cs_pyecharts\out1.png', width=Inches(1.0))
# document.add_paragraph(text='\r', style=None)
# p = document.add_paragraph()
# run = p.add_run('音乐原声')
# run.font.size = Pt(22)  # 二号
# run.bold = True


table = document.add_table(rows=14, cols=6)  # 4行5列的表格
# table.add_column(width=Cm(1))
from docx.enum.table import WD_TABLE_ALIGNMENT
table.alignment = WD_TABLE_ALIGNMENT.CENTER #表格居中
table.style = document.styles['Table Grid']
table.autofit=True
def th(x,y,content):
    """
    th样式
    :param x: x坐标
    :param y: y坐标
    :param content: 内容
    :return: None
    """
    # print(grid,content)
    run = table.cell(x,y).paragraphs[0].add_run(content)
    run.bold = True  # 加粗

#table.cell(0,0).text = "功能测试报告"
th(0,0,"功能测试报告")
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
for x in range(14):
    table.rows[x].height=Cm(1)
    for y in range(6):
        table.cell(x,y).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
table.cell(0,1).text = "自助打印"
# table.cell(0,2).text = "作曲"
# table.cell(0,3).text = "作词"
# table.cell(0,4).text = "类型"
num=1
for i in range(num + 1, 6):
    table.cell(0, num).merge(table.cell(0, i))


table.cell(1,0).text = "测试版本"
table.cell(1,1).text = "1.1.0.34"
table.cell(1, 1).merge(table.cell(1, 2))
table.cell(1,3).text = "测试人员"
table.cell(1,4).text = "鲁潜元"
table.cell(1,4).merge(table.cell(1,5))


table.cell(2,0).text = "测试时间"
table.cell(2,1).text = "2020-07-28"
table.cell(2,2).text = "2020-07-28"
table.cell(2,3).text = "测试耗时"
table.cell(2,4).text = "1MH"
table.cell(2,4).merge(table.cell(2,5))

table.cell(3,0).merge(table.cell(4,0))
table.cell(3,0).text = "测试环境"
table.cell(3,1).text = "url"
table.cell(3,2).text='url地址'
for i in range(2, 6):
    table.cell(3, 3).merge(table.cell(3, i))

table.cell(4,1).text = "client"
table.cell(4,2).text='环境地址'
for i in range(2, 6):
    table.cell(4, 3).merge(table.cell(4, i))
document.save('cs.docx')

table.cell(5,0).text = "测试工具"
table.cell(5,1).text='测试工具地址'
for i in range(1, 6):
    table.cell(5, 1).merge(table.cell(5, i))

table.cell(6,0).text = "测试内容"
table.cell(6,1).text='测试内容描述'
for i in range(1, 6):
    table.cell(6, 1).merge(table.cell(6, i))

table.cell(7,0).merge(table.cell(8,0))
table.cell(7,0).text = "BUG说明"
table.cell(7,1).text = """【新增情况】
级别分布 
类型分布"""
for i in range(1, 6):
    table.cell(7, 1).merge(table.cell(7, i))

table.cell(8,1).text =f"""
【回归情况】 
总计 2
关闭 2
激活 0 """
table.cell(8,2).text = "趋势统计图"
for i in range(2, 6):
    table.cell(8, 2).merge(table.cell(8, i))

table.cell(9,0).text = "CASE执行"
table.cell(9,1).text='CASE执行情况'
for i in range(1, 6):
    table.cell(9, 1).merge(table.cell(9, i))

table.cell(10,0).text = "风险评估"
table.cell(10,1).text='风险评估'
for i in range(1, 6):
    table.cell(10, 1).merge(table.cell(10, i))

table.cell(11,0).text = "测试总结"
table.cell(11,1).text='''1.bug回归测试
2.本轮测试暂未发现明显错误，BUG已修复23333333333333333333对该版本进行发布'''
for i in range(1, 6):
    table.cell(11, 1).merge(table.cell(11, i))

table.cell(12,0).text = "IN值"
table.cell(12,0).width=Inches(1.1)
table.cell(12,1).text = "IN计算"
table.cell(12,2).text = "提测质量"
table.cell(12,3).text = "合格"
table.cell(12,4).text = "发布评估"
table.cell(12,5).text = "发布"

table.cell(13,0).text = "报告连接"
table.cell(13,1).text='报告连接地址'
for i in range(1, 6):
    table.cell(13, 1).merge(table.cell(13, i))

document.save('cs.docx')